from io import BytesIO

import pandas as pd
from docx import Document

from utils.word_report import build_calculation_report_docx


def test_word_report_readability_sections_and_missing_values():
    data = pd.DataFrame(
        {
            "管段编号": ["S1-1", "S1-2"],
            "风量": [1000, None],
            "校核": ["合理", pd.NA],
        }
    )

    docx_bytes = build_calculation_report_docx(
        title="测试计算说明书",
        module_name="测试模块",
        description="用于测试 Word 导出可读性。",
        input_tables=[{"title": "输入参数表", "data": data}],
        result_tables=[{"title": "计算结果表", "data": data}],
        summary_rows=[("系统总阻力", "100 Pa")],
        formula_rows=[("测试公式", "Q = q / 3600", "Q：m³/s")],
        notes=["测试说明，不能替代正式工程设计。"],
    )

    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 1000

    doc = Document(BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    full_text = text + "\n" + table_text

    assert "测试计算说明书" in full_text
    assert "测试模块" in full_text
    assert "输入参数表" in full_text
    assert "计算结果表" in full_text
    assert "系统总阻力" in full_text
    assert "测试公式" in full_text
    assert "说明与免责声明" in full_text
    assert "当前稳定版本：v0.2.0" in full_text
    assert "—" in full_text
    assert "NaN" not in full_text
    assert "<NA>" not in full_text
    assert "正式工程可直接使用" not in full_text
