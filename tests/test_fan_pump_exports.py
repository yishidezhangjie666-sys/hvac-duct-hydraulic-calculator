from io import BytesIO
from pathlib import Path
import sys

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


def _assert_docx_contains(data, expected_title):
    assert isinstance(data, bytes)
    assert len(data) > 1000

    doc = Document(BytesIO(data))
    paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    full_text = paragraph_text + "\n" + table_text
    assert expected_title in paragraph_text
    assert "说明与免责声明" in paragraph_text
    assert len(doc.tables) >= 3
    assert "None" not in full_text
    assert "NaN" not in full_text
    assert "<NA>" not in full_text


def test_fan_word_report_can_be_opened():
    from modules.fan_pump_selection import (
        FAN_EXPORT_MAP,
        FAN_SAMPLE_DATA,
        _blank_missing_export_values,
        _build_fan_word_report,
        _format_number_columns,
        _format_percent_columns,
        _rename_export_columns,
        _summary_rows,
        calculate_fan_selection,
    )

    result = calculate_fan_selection(FAN_SAMPLE_DATA)
    df_export = _blank_missing_export_values(
        _format_number_columns(
            _format_percent_columns(
                _rename_export_columns(result, FAN_EXPORT_MAP),
                ["风量余量", "风压余量", "电机功率余量"],
            ),
            ["风机轴功率参考（kW）"],
        )
    )

    data = _build_fan_word_report(df_export, _summary_rows(result))
    _assert_docx_contains(data, "风机选型校核计算说明书")


def test_pump_word_report_can_be_opened():
    from modules.fan_pump_selection import (
        PUMP_EXPORT_MAP,
        PUMP_SAMPLE_DATA,
        _blank_missing_export_values,
        _build_pump_word_report,
        _format_number_columns,
        _format_percent_columns,
        _rename_export_columns,
        _summary_rows,
        calculate_pump_selection,
    )

    result = calculate_pump_selection(PUMP_SAMPLE_DATA)
    df_export = _blank_missing_export_values(
        _format_number_columns(
            _format_percent_columns(
                _rename_export_columns(result, PUMP_EXPORT_MAP),
                ["流量余量", "扬程余量", "电机功率余量"],
            ),
            ["水泵轴功率参考（kW）"],
        )
    )

    data = _build_pump_word_report(df_export, _summary_rows(result))
    _assert_docx_contains(data, "水泵选型校核计算说明书")
