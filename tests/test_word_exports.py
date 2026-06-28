from io import BytesIO
from pathlib import Path
import sys

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


def _assert_valid_docx(data, title_part):
    assert isinstance(data, bytes)
    assert len(data) > 1000

    doc = Document(BytesIO(data))
    texts = [paragraph.text for paragraph in doc.paragraphs]
    assert any(title_part in text for text in texts)
    assert any("说明与免责声明" in text for text in texts)
    assert len(doc.tables) >= 1


def test_build_calculation_report_docx_basic():
    from utils.word_report import build_calculation_report_docx

    data = build_calculation_report_docx(
        title="测试计算说明书",
        module_name="测试模块",
        description="用于验证统一 Word 报告结构。",
        input_tables=[
            {"title": "输入数据", "data": pd.DataFrame([{"参数": "Q", "数值": "1000 m³/h"}])}
        ],
        result_tables=[
            {"title": "计算结果", "data": [{"项目": "系统总阻力", "数值": "60 Pa"}]}
        ],
        summary_rows=[("方案数量", "1 个")],
        formula_rows=[("示例公式", "Q = Q_h / 3600", "Q：m³/s")],
    )

    _assert_valid_docx(data, "测试计算说明书")


def test_ventilation_word_export_uses_unified_structure():
    from modules import ventilation_duct

    df_export = pd.DataFrame(
        [
            {
                "管段编号": "S1-1",
                "风量 Q（m³/h）": 1200,
                "宽度 a（mm）": 500,
                "高度 b（mm）": 320,
                "长度 L（m）": 8,
                "单位长度摩擦阻力 R（Pa/m）": 1.2,
                "局部阻力系数 ζ": 0.8,
                "风量 q（m³/s）": 0.333,
                "截面积 A（m²）": 0.16,
                "风速 v（m/s）": 2.08,
                "当量直径 De（m）": 0.39,
                "动压 Pd（Pa）": 2.6,
                "沿程阻力 Py（Pa）": 9.6,
                "局部阻力 Pj（Pa）": 2.08,
                "管段总阻力 Pi（Pa）": 11.68,
            }
        ]
    )
    summary_rows = [
        ("系统总风量", "1200.00 m³/h"),
        ("系统总阻力", "11.68 Pa"),
        ("推荐风机风量", "1200.00 m³/h"),
        ("推荐风机风压", "11.68 Pa"),
    ]

    data = ventilation_duct._build_ventilation_word_report(
        df_export,
        total_airflow=1200,
        system_total=11.68,
        rho=1.2,
        summary_rows=summary_rows,
    )

    _assert_valid_docx(data, "通风风管水力计算说明书")


def test_water_system_word_export_uses_unified_structure():
    from modules import air_conditioning_water as water
    from utils.export_utils import WATER_EXPORT_MAP, rename_export_columns

    df = pd.DataFrame(water.SAMPLE_DATA)
    df = water.calculate_flows(df, delta_t=5.0)
    result, summary = water.calculate_water_system(
        df,
        rho=1000,
        flow_safety_factor=1.10,
        pressure_safety_factor=1.15,
    )
    df_export = rename_export_columns(result, WATER_EXPORT_MAP)
    summary_rows = [
        ("系统总流量", f"{summary['total_flow']:.2f} m³/h"),
        ("系统总阻力", f"{summary['total_loss']:.1f} Pa"),
        ("推荐水泵流量", f"{summary['pump_flow']:.2f} m³/h"),
        ("推荐水泵扬程", f"{summary['pump_head_m']:.2f} m"),
        ("推荐水泵扬程", f"{summary['pump_head_kpa']:.1f} kPa"),
    ]

    data = water._build_water_word_report(
        df_export,
        summary,
        rho=1000,
        cp=4.186,
        delta_t=5.0,
        flow_safety_factor=1.10,
        pressure_safety_factor=1.15,
        summary_rows=summary_rows,
    )

    _assert_valid_docx(data, "空调水系统水力计算说明书")


def test_terminal_equipment_word_exports_use_unified_structure():
    from modules import terminal_equipment as terminal

    fcu_result = terminal.calculate_fcu_selection(terminal.FCU_SAMPLE_DATA)
    fcu_export = terminal._format_percent_columns(
        terminal._rename_export_columns(fcu_result, terminal.FCU_EXPORT_MAP),
        ["冷量余量", "热量余量", "风量余量"],
    )
    fcu_doc = terminal._build_terminal_word_report(
        "风机盘管初步选型计算说明书",
        "风机盘管初步选型",
        "用于风机盘管冷量、热量和风量的简化初步选型校核。",
        fcu_export,
        terminal._fcu_summary_rows(fcu_result),
        terminal.FCU_INPUT_COLUMNS,
        terminal.FCU_RESULT_COLUMNS,
        terminal.FCU_FORMULA_ROWS,
    )

    pau_result = terminal.calculate_pau_selection(terminal.PAU_SAMPLE_DATA)
    pau_export = terminal._format_percent_columns(
        terminal._rename_export_columns(pau_result, terminal.PAU_EXPORT_MAP),
        ["冷量余量", "风量余量"],
    )
    pau_doc = terminal._build_terminal_word_report(
        "新风机组初步选型计算说明书",
        "新风机组初步选型",
        "用于新风机组冷量和风量的简化初步选型校核。",
        pau_export,
        terminal._pau_summary_rows(pau_result),
        terminal.PAU_INPUT_COLUMNS,
        terminal.PAU_RESULT_COLUMNS,
        terminal.PAU_FORMULA_ROWS,
    )

    _assert_valid_docx(fcu_doc, "风机盘管初步选型计算说明书")
    _assert_valid_docx(pau_doc, "新风机组初步选型计算说明书")


def test_heat_cold_source_word_exports_use_unified_structure():
    from modules import heat_cold_source as source

    air_result = source.calculate_air_heat_pump_selection(source.AIR_HEAT_PUMP_SAMPLE_DATA)
    air_export = source._format_percent_columns(
        source._rename_export_columns(air_result, source.AIR_HEAT_PUMP_EXPORT_MAP),
        ["制冷量余量", "制热量余量"],
    )
    air_doc = source._build_source_word_report(
        "风冷热泵机组初步选型计算说明书",
        "风冷热泵机组初步选型",
        "用于风冷热泵机组制冷量、制热量和台数的简化初步选型校核。",
        air_export,
        source._summary_rows(air_result, "scheme_id", "总制冷量 (kW)"),
        source.AIR_HEAT_PUMP_INPUT_COLUMNS,
        source.AIR_HEAT_PUMP_RESULT_COLUMNS,
        source.SOURCE_FORMULA_ROWS,
    )

    chiller_result = source.calculate_chiller_selection(source.CHILLER_SAMPLE_DATA)
    chiller_export = source._format_number_columns(
        source._format_percent_columns(
            source._rename_export_columns(chiller_result, source.CHILLER_EXPORT_MAP),
            ["制冷量余量"],
        ),
        ["供回水温差 Δt（℃）"],
    )
    chiller_doc = source._build_source_word_report(
        "冷水机组初步选型计算说明书",
        "冷水机组初步选型",
        "用于冷水机组制冷量、台数和冷冻水温差的简化初步选型校核。",
        chiller_export,
        source._summary_rows(chiller_result, "scheme_id", "总制冷量 (kW)"),
        source.CHILLER_INPUT_COLUMNS,
        source.CHILLER_RESULT_COLUMNS,
        source.SOURCE_FORMULA_ROWS + [source.COOLING_DELTA_FORMULA_ROW],
    )

    boiler_result = source.calculate_boiler_selection(source.BOILER_SAMPLE_DATA)
    boiler_export = source._format_number_columns(
        source._format_percent_columns(
            source._rename_export_columns(boiler_result, source.BOILER_EXPORT_MAP),
            ["供热量余量"],
        ),
        ["供回水温差 Δt（℃）"],
    )
    boiler_doc = source._build_source_word_report(
        "锅炉 / 热源设备初步选型计算说明书",
        "锅炉 / 热源设备初步选型",
        "用于锅炉或热源设备供热量、台数和供回水温差的简化初步选型校核。",
        boiler_export,
        source._summary_rows(boiler_result, "scheme_id", "总供热量 (kW)"),
        source.BOILER_INPUT_COLUMNS,
        source.BOILER_RESULT_COLUMNS,
        source.SOURCE_FORMULA_ROWS + [source.HEATING_DELTA_FORMULA_ROW],
    )

    _assert_valid_docx(air_doc, "风冷热泵机组初步选型计算说明书")
    _assert_valid_docx(chiller_doc, "冷水机组初步选型计算说明书")
    _assert_valid_docx(boiler_doc, "锅炉 / 热源设备初步选型计算说明书")
