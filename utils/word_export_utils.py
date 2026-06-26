"""
Word 计算说明书导出工具函数

依赖 python-docx，用于生成格式化的 .docx 计算说明书。
"""

import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── 字体辅助函数 ──────────────────────────────────

def _set_run(run, size=11, bold=False, italic=False, subscript=False, east_asia=None):
    """设置 run 的完整格式：西文 Times New Roman，中文 east_asia，黑色"""
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    run.font.subscript = subscript
    run.font.name = "Times New Roman"
    if east_asia:
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), east_asia)


def _add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        _set_run(run, size=16, bold=True, east_asia="黑体")


def _add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        _set_run(run, size=13, bold=True, east_asia="黑体")


def _add_para(doc, text):
    p = doc.add_paragraph()
    _set_run(p.add_run(text), east_asia="宋体")


def _add_table(doc, headers, rows, table_title=None):
    if table_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(p.add_run(table_title), size=10, bold=True, east_asia="黑体")

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 等宽列
    pw = int(Cm(15.5)) // len(headers)
    for row_obj in table.rows:
        for cell in row_obj.cells:
            cell.width = pw

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        _set_run(para.add_run(str(h)), size=9, bold=True, east_asia="黑体")

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            _set_run(para.add_run(str(val)), size=9, east_asia="宋体")


def _add_disclaimer(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.italic = True


# ─── 富文本段落（真实下标） ────────────────────────

def _rich(doc, parts):
    p = doc.add_paragraph()
    for text, italic, sub in parts:
        _set_run(p.add_run(text), italic=italic, subscript=sub, east_asia="宋体")
    return p


# 简写
_N = lambda t: (t, False, False)
_I = lambda t: (t, True, False)
_S = lambda t: (t, True, True)


# ─── 数值格式化 ────────────────────────────────────

def _round_df(df, fmt):
    df = df.copy()
    for col, dec in fmt.items():
        if col in df.columns:
            df[col] = df[col].apply(
                lambda v, d=dec: round(v, d) if isinstance(v, (int, float)) else v
            )
    return df


_VENT_FMT = {
    "风量 Q（m³/h）": 0, "宽度 a（mm）": 0, "高度 b（mm）": 0,
    "长度 L（m）": 1, "单位长度摩擦阻力 R（Pa/m）": 2, "局部阻力系数 ζ": 2,
    "风量 q（m³/s）": 3, "截面积 A（m²）": 3, "风速 v（m/s）": 2,
    "当量直径 De（m）": 3, "动压 Pd（Pa）": 1, "沿程阻力 Py（Pa）": 1,
    "局部阻力 Pj（Pa）": 1, "管段总阻力 Pi（Pa）": 1,
}

_WATER_FMT = {
    "负荷 QL（kW）": 1, "水流量 G（m³/h）": 2, "管道内径 D（mm）": 0,
    "长度 L（m）": 1, "单位长度沿程阻力 R（Pa/m）": 1, "局部阻力系数 ζ": 1,
    "流量 q（m³/s）": 4, "截面积 A（m²）": 4, "流速 v（m/s）": 2,
    "动压 Pd（Pa）": 1, "沿程阻力 Py（Pa）": 1, "局部阻力 Pj（Pa）": 1,
    "管段总阻力 Pi（Pa）": 1,
}


# ─── 表格拆分列名 ──────────────────────────────────
# 注意：以下列名必须与 export_utils.py 中 VENTILATION_EXPORT_MAP /
# WATER_EXPORT_MAP 的 values 完全一致。

_VENT_INPUT = [
    "管段编号", "风量 Q（m³/h）", "宽度 a（mm）", "高度 b（mm）",
    "长度 L（m）", "单位长度摩擦阻力 R（Pa/m）", "局部阻力系数 ζ",
]
_VENT_RESULT = [
    "管段编号", "风量 q（m³/s）", "截面积 A（m²）", "风速 v（m/s）",
    "当量直径 De（m）", "动压 Pd（Pa）", "沿程阻力 Py（Pa）",
    "局部阻力 Pj（Pa）", "管段总阻力 Pi（Pa）",
]

_WATER_INPUT = [
    "管段编号", "负荷 QL（kW）", "水流量 G（m³/h）", "管道内径 D（mm）",
    "长度 L（m）", "单位长度沿程阻力 R（Pa/m）", "局部阻力系数 ζ",
]
_WATER_RESULT = [
    "管段编号", "流量 q（m³/s）", "截面积 A（m²）", "流速 v（m/s）",
    "动压 Pd（Pa）", "沿程阻力 Py（Pa）", "局部阻力 Pj（Pa）",
    "管段总阻力 Pi（Pa）", "流速校核",
]


# ─── 公式定义（Word 真实下标） ────────────────────

_VENT_FMLS = [
    [_N("风量换算："), _I("Q"), _N(" = "), _I("Q"), _S("h"), _N(" / 3600（"), _I("Q"), _N("：m³/s，"), _I("Q"), _S("h"), _N("：m³/h）")],
    [_N("截面积："), _I("A"), _N(" = "), _I("a"), _N(" × "), _I("b"), _N("（"), _I("A"), _N("：m²，"), _I("a"), _N("、"), _I("b"), _N("：m）")],
    [_N("风速："), _I("v"), _N(" = "), _I("Q"), _N(" / "), _I("A"), _N("（"), _I("v"), _N("：m/s）")],
    [_N("水力直径："), _I("D"), _S("e"), _N(" = 2"), _I("ab"), _N(" / ("), _I("a"), _N(" + "), _I("b"), _N(")（"), _I("D"), _S("e"), _N("：m）")],
    [_N("动压："), _I("P"), _S("d"), _N(" = ρ"), _I("v"), _N("² / 2（"), _I("P"), _S("d"), _N("：Pa，ρ：kg/m³）")],
    [_N("沿程阻力："), _I("P"), _S("y"), _N(" = "), _I("R"), _N(" × "), _I("L"), _N("（"), _I("P"), _S("y"), _N("：Pa，"), _I("R"), _N("：Pa/m，"), _I("L"), _N("：m）")],
    [_N("局部阻力："), _I("P"), _S("j"), _N(" = ζ × "), _I("P"), _S("d"), _N("（"), _I("P"), _S("j"), _N("：Pa）")],
    [_N("管段总阻力："), _I("P"), _S("i"), _N(" = "), _I("P"), _S("y"), _N(" + "), _I("P"), _S("j"), _N("（"), _I("P"), _S("i"), _N("：Pa）")],
    [_N("系统总阻力：Σ"), _I("P"), _N(" = Σ"), _I("P"), _S("i"), _N("（Pa）")],
]

_WATER_FMLS = [
    [_N("水流量估算："), _I("G"), _N(" = 0.86 × "), _I("Q"), _S("L"), _N(" / Δ"), _I("t"), _N("（"), _I("G"), _N("：m³/h，"), _I("Q"), _S("L"), _N("：kW，Δ"), _I("t"), _N("：℃）")],
    [_N("流量换算："), _I("q"), _N(" = "), _I("G"), _N(" / 3600（"), _I("q"), _N("：m³/s）")],
    [_N("截面积："), _I("A"), _N(" = π"), _I("D"), _N("² / 4（"), _I("A"), _N("：m²，"), _I("D"), _N("：m）")],
    [_N("流速："), _I("v"), _N(" = "), _I("q"), _N(" / "), _I("A"), _N("（"), _I("v"), _N("：m/s）")],
    [_N("动压："), _I("P"), _S("d"), _N(" = ρ"), _I("v"), _N("² / 2（"), _I("P"), _S("d"), _N("：Pa，ρ：kg/m³）")],
    [_N("沿程阻力："), _I("P"), _S("y"), _N(" = "), _I("R"), _N(" × "), _I("L"), _N("（"), _I("P"), _S("y"), _N("：Pa，"), _I("R"), _N("：Pa/m，"), _I("L"), _N("：m）")],
    [_N("局部阻力："), _I("P"), _S("j"), _N(" = ζ × "), _I("P"), _S("d"), _N("（"), _I("P"), _S("j"), _N("：Pa）")],
    [_N("管段总阻力："), _I("P"), _S("i"), _N(" = "), _I("P"), _S("y"), _N(" + "), _I("P"), _S("j"), _N("（"), _I("P"), _S("i"), _N("：Pa）")],
    [_N("系统总阻力：Σ"), _I("P"), _N(" = Σ"), _I("P"), _S("i"), _N("（Pa）")],
    [_N("水泵流量："), _I("G"), _S("pump"), _N(" = "), _I("G"), _S("total"), _N(" × "), _I("k"), _S("f"), _N("（"), _I("k"), _S("f"), _N("：流量安全系数）")],
    [_N("水泵扬程："), _I("H"), _N(" = Σ"), _I("P"), _N(" × "), _I("k"), _S("p"), _N(" / (ρ"), _I("g"), _N(")（"), _I("k"), _S("p"), _N("：扬程安全系数，"), _I("g"), _N("：9.81 m/s²）")],
]


# ─── 取子表 ────────────────────────────────────────

def _subtable(df, cols):
    available = [c for c in cols if c in df.columns]
    return df[available].copy()


# ─── 通风模块 Word 报告 ────────────────────────────

def build_ventilation_word_report(df_export, total_airflow, system_total, rho):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_title(doc, "通风风管水力计算说明书")
    _add_para(doc, "本说明书由「建环工程计算工具箱」自动生成，用于通风风管系统课程设计和工程初步校核。")

    _add_h1(doc, "一、输入参数")
    _add_para(doc, f"空气密度 ρ = {rho:.2f} kg/m³")
    _add_para(doc, "风机风量安全系数 = 1.0（当前版本暂未设置独立安全系数参数）")
    _add_para(doc, "风机风压安全系数 = 1.0（当前版本暂未设置独立安全系数参数）")

    _add_h1(doc, "二、管段计算结果")
    df_fmt = _round_df(df_export, _VENT_FMT)
    _add_table(doc, _VENT_INPUT, _subtable(df_fmt, _VENT_INPUT).values.tolist(), "表 1  管段输入参数表")
    doc.add_paragraph()
    _add_table(doc, _VENT_RESULT, _subtable(df_fmt, _VENT_RESULT).values.tolist(), "表 2  管段计算结果表")

    _add_h1(doc, "三、系统汇总")
    _add_para(doc, f"系统总风量：{total_airflow:.2f} m³/h")
    _add_para(doc, f"系统总阻力：{system_total:.2f} Pa")
    _add_para(doc, f"推荐风机风量：{total_airflow:.2f} m³/h（安全系数 1.0）")
    _add_para(doc, f"推荐风机风压：{system_total:.2f} Pa（安全系数 1.0）")

    _add_h1(doc, "四、计算公式说明")
    for parts in _VENT_FMLS:
        _rich(doc, parts)

    _add_h1(doc, "五、免责声明")
    _add_disclaimer(doc, "计算结果仅供课程设计和工程初步校核使用，实际工程应结合规范、设备样本及最不利环路进行复核。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 空调水系统模块 Word 报告 ──────────────────────

def build_water_system_word_report(df_export, summary, rho, cp, delta_t, flow_safety_factor, pressure_safety_factor):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _add_title(doc, "空调水系统水力计算说明书")
    _add_para(doc, "本说明书由「建环工程计算工具箱」自动生成，用于空调冷冻水、热水系统课程设计和工程初步校核。")

    _add_h1(doc, "一、输入参数")
    _add_para(doc, f"水密度 ρ = {rho:.0f} kg/m³")
    _add_para(doc, f"水比热容 c = {cp:.3f} kJ/(kg·℃)")
    _add_para(doc, f"供回水温差 Δt = {delta_t:.1f} ℃")
    _add_para(doc, f"水泵流量安全系数 = {flow_safety_factor:.2f}")
    _add_para(doc, f"水泵扬程安全系数 = {pressure_safety_factor:.2f}")

    _add_h1(doc, "二、管段计算结果")
    df_fmt = _round_df(df_export, _WATER_FMT)
    _add_table(doc, _WATER_INPUT, _subtable(df_fmt, _WATER_INPUT).values.tolist(), "表 1  管段输入参数表")
    doc.add_paragraph()
    _add_table(doc, _WATER_RESULT, _subtable(df_fmt, _WATER_RESULT).values.tolist(), "表 2  管段计算结果表")

    _add_h1(doc, "三、系统汇总")
    _add_para(doc, f"系统总流量：{summary['total_flow']:.2f} m³/h")
    _add_para(doc, f"系统总阻力：{summary['total_loss']:.1f} Pa")
    _add_para(doc, f"推荐水泵流量：{summary['pump_flow']:.2f} m³/h")
    _add_para(doc, f"推荐水泵扬程：{summary['pump_head_m']:.2f} m")
    _add_para(doc, f"推荐水泵扬程：{summary['pump_head_kpa']:.1f} kPa")

    _add_h1(doc, "四、计算公式说明")
    for parts in _WATER_FMLS:
        _rich(doc, parts)

    _add_h1(doc, "五、免责声明")
    _add_disclaimer(doc, "计算结果仅供课程设计和工程初步校核使用，实际工程应结合规范、设备样本、最不利环路及水力平衡要求复核。")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
