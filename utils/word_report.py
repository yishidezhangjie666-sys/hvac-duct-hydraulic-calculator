"""Unified Word calculation report builder."""

from datetime import datetime
import io

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt


DEFAULT_NOTES = [
    "本工具采用简化工程计算口径，结果仅用于学习、课程设计辅助核算和工程初步校核，不能替代正式工程设计、设备样本选型或规范校核。",
]


def _safe_text(value, missing="—"):
    """Return a stable text value for Word cells and paragraphs."""
    if value is None:
        return missing
    try:
        if pd.isna(value):
            return missing
    except (TypeError, ValueError):
        pass
    return str(value)


def _configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Title"]:
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Microsoft YaHei"


def _set_paragraph_font(paragraph, size=10.5, bold=None):
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def _set_cell_text(cell, text, bold=False):
    cell.text = _safe_text(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        _set_paragraph_font(paragraph, size=9, bold=bold)


def _add_heading(doc, text, level):
    doc.add_heading(_safe_text(text), level=level)


def _add_paragraphs(doc, paragraphs):
    for paragraph in paragraphs or []:
        text = _safe_text(paragraph, missing="")
        if text:
            new_paragraph = doc.add_paragraph(text)
            _set_paragraph_font(new_paragraph)


def _normalize_table_data(table_data):
    if isinstance(table_data, pd.DataFrame):
        headers = [_safe_text(col) for col in table_data.columns]
        rows = [
            [_safe_text(value) for value in row]
            for row in table_data.itertuples(index=False, name=None)
        ]
        return headers, rows

    if table_data is None:
        return [], []

    rows_source = list(table_data)
    if not rows_source:
        return [], []

    if all(isinstance(row, dict) for row in rows_source):
        headers = []
        for row in rows_source:
            for key in row.keys():
                if key not in headers:
                    headers.append(key)
        rows = [[_safe_text(row.get(header, "")) for header in headers] for row in rows_source]
        return [_safe_text(header) for header in headers], rows

    rows_as_lists = []
    for row in rows_source:
        if isinstance(row, (list, tuple)):
            rows_as_lists.append([_safe_text(value) for value in row])
        else:
            rows_as_lists.append([_safe_text(row)])

    max_cols = max(len(row) for row in rows_as_lists)
    if max_cols == 2:
        headers = ["项目", "数值"]
    else:
        headers = [f"列 {idx}" for idx in range(1, max_cols + 1)]
    rows = [row + [""] * (max_cols - len(row)) for row in rows_as_lists]
    return headers, rows


def _coerce_table_entry(entry, default_title):
    if isinstance(entry, dict):
        return entry.get("title", default_title), entry.get("data", [])
    if isinstance(entry, (list, tuple)) and len(entry) == 2 and isinstance(entry[0], str):
        return entry[0], entry[1]
    return default_title, entry


def _add_table(doc, title, df_or_rows):
    title_text = _safe_text(title)
    if title_text:
        title_paragraph = doc.add_paragraph(title_text)
        _set_paragraph_font(title_paragraph, bold=True)

    headers, rows = _normalize_table_data(df_or_rows)
    if not headers or not rows:
        doc.add_paragraph("暂无数据")
        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[idx], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            _set_cell_text(cells[idx], value)

    doc.add_paragraph("")


def _add_key_value_rows(doc, rows):
    normalized_rows = []
    for row in rows or []:
        if isinstance(row, dict):
            normalized_rows.append(
                {
                    "项目": row.get("项目", row.get("label", "")),
                    "数值": row.get("数值", row.get("value", "")),
                }
            )
        elif isinstance(row, (list, tuple)) and len(row) >= 2:
            normalized_rows.append({"项目": row[0], "数值": row[1]})
        else:
            normalized_rows.append({"项目": row, "数值": ""})

    _add_table(doc, "系统汇总表", normalized_rows)


def _normalize_formula_rows(formula_rows):
    normalized_rows = []
    for row in formula_rows or []:
        if isinstance(row, dict):
            normalized_rows.append(
                {
                    "公式名称": row.get("公式名称", row.get("name", "")),
                    "公式表达": row.get("公式表达", row.get("formula", "")),
                    "单位说明": row.get("单位说明", row.get("unit", "")),
                }
            )
        elif isinstance(row, (list, tuple)):
            values = list(row) + ["", "", ""]
            normalized_rows.append(
                {
                    "公式名称": values[0],
                    "公式表达": values[1],
                    "单位说明": values[2],
                }
            )
        else:
            normalized_rows.append({"公式名称": row, "公式表达": "", "单位说明": ""})
    return normalized_rows


def build_calculation_report_docx(
    title,
    module_name,
    description,
    input_tables,
    result_tables,
    summary_rows,
    formula_rows,
    notes=None,
):
    """Build a unified calculation report and return .docx bytes."""
    doc = Document()
    _configure_document(doc)

    _add_heading(doc, title, 0)
    _add_paragraphs(
        doc,
        [
            "生成工具：建环工程计算工具箱",
            "当前稳定版本：v0.2.0",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        ],
    )

    _add_heading(doc, "一、模块说明", 1)
    _add_paragraphs(
        doc,
        [
            f"模块名称：{_safe_text(module_name)}",
            f"功能说明：{_safe_text(description)}",
            "适用范围：适用于学习、课程设计辅助核算、工程初步校核和个人作品集展示。",
            "工程边界提示：本说明书为普通计算说明书，不作为正式工程设计文件。",
        ],
    )

    _add_heading(doc, "二、输入参数", 1)
    for idx, table_data in enumerate(input_tables or [], start=1):
        table_title, data = _coerce_table_entry(table_data, f"输入参数表 {idx}")
        _add_table(doc, table_title, data)
    if not input_tables:
        doc.add_paragraph("暂无数据")

    _add_heading(doc, "三、计算结果", 1)
    for idx, table_data in enumerate(result_tables or [], start=1):
        table_title, data = _coerce_table_entry(table_data, f"计算结果表 {idx}")
        _add_table(doc, table_title, data)
    if not result_tables:
        doc.add_paragraph("暂无数据")

    _add_heading(doc, "四、系统汇总", 1)
    _add_key_value_rows(doc, summary_rows)

    _add_heading(doc, "五、公式说明", 1)
    _add_table(doc, "公式说明表", _normalize_formula_rows(formula_rows))

    _add_heading(doc, "六、说明与免责声明", 1)
    _add_paragraphs(doc, (notes or []) + DEFAULT_NOTES)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
