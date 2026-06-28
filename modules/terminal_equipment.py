"""
空调末端设备初步选型模块

模块三：风机盘管和新风机组的简化初步选型校核。
"""

import io

import pandas as pd
import streamlit as st
from docx import Document

from utils.export_utils import export_formatted_excel, get_csv_bytes


FCU_SAMPLE_DATA = {
    "room_id": ["R-101", "R-102", "R-201", "R-202"],
    "room_name": ["办公室", "会议室", "资料室", "开放办公区"],
    "cooling_load_kw": [3.2, 5.8, 2.4, 8.5],
    "heating_load_kw": [2.8, 4.6, 2.0, 7.2],
    "design_airflow_m3h": [520, 860, 420, 1250],
    "rated_cooling_kw": [3.8, 6.3, 2.8, 10.0],
    "rated_heating_kw": [3.4, 5.2, 2.5, 8.4],
    "rated_airflow_m3h": [600, 900, 450, 1400],
}

PAU_SAMPLE_DATA = {
    "system_id": ["PAU-1", "PAU-2"],
    "service_area": ["一层办公区", "二层会议区"],
    "fresh_airflow_m3h": [2500, 1800],
    "outdoor_enthalpy_kjkg": [85.0, 82.0],
    "indoor_enthalpy_kjkg": [55.0, 54.0],
    "air_density_kgm3": [1.20, 1.20],
    "rated_cooling_kw": [30.0, 18.0],
    "rated_airflow_m3h": [2800, 2000],
}

FCU_EXPORT_MAP = {
    "room_id": "房间编号",
    "room_name": "房间名称",
    "cooling_load_kw": "冷负荷 Qc（kW）",
    "heating_load_kw": "热负荷 Qh（kW）",
    "design_airflow_m3h": "设计风量 L（m³/h）",
    "rated_cooling_kw": "额定冷量（kW）",
    "rated_heating_kw": "额定热量（kW）",
    "rated_airflow_m3h": "额定风量（m³/h）",
    "冷量余量": "冷量余量",
    "热量余量": "热量余量",
    "风量余量": "风量余量",
    "冷量校核": "冷量校核",
    "热量校核": "热量校核",
    "风量校核": "风量校核",
    "选型结论": "选型结论",
}

PAU_EXPORT_MAP = {
    "system_id": "系统编号",
    "service_area": "服务区域",
    "fresh_airflow_m3h": "新风量 L（m³/h）",
    "outdoor_enthalpy_kjkg": "室外空气焓值 hW（kJ/kg）",
    "indoor_enthalpy_kjkg": "室内空气焓值 hN（kJ/kg）",
    "air_density_kgm3": "空气密度 ρ（kg/m³）",
    "rated_cooling_kw": "拟选新风机组冷量（kW）",
    "rated_airflow_m3h": "拟选新风机组风量（m³/h）",
    "新风质量流量 m (kg/s)": "新风质量流量 m（kg/s）",
    "新风冷负荷 Q (kW)": "新风冷负荷 Q（kW）",
    "冷量余量": "冷量余量",
    "风量余量": "风量余量",
    "冷量校核": "冷量校核",
    "风量校核": "风量校核",
    "焓差校核": "焓差校核",
    "选型结论": "选型结论",
}

FORMULA_STYLE = """
<style>
.formula-table {
    width: 100%;
    border-collapse: collapse;
}
.formula-table th, .formula-table td {
    border: 1px solid rgba(255,255,255,0.18);
    padding: 10px 14px;
    text-align: left;
}
.formula-table th {
    font-weight: 700;
}
.formula-table sub {
    vertical-align: sub;
    font-size: 75%;
    line-height: 0;
}
.formula-table sup {
    vertical-align: super;
    font-size: 75%;
    line-height: 0;
}
</style>
"""


def _empty_fcu_df():
    return pd.DataFrame(
        {
            "room_id": pd.Series(dtype="str"),
            "room_name": pd.Series(dtype="str"),
            "cooling_load_kw": pd.Series(dtype="float64"),
            "heating_load_kw": pd.Series(dtype="float64"),
            "design_airflow_m3h": pd.Series(dtype="float64"),
            "rated_cooling_kw": pd.Series(dtype="float64"),
            "rated_heating_kw": pd.Series(dtype="float64"),
            "rated_airflow_m3h": pd.Series(dtype="float64"),
        }
    )


def _empty_pau_df():
    return pd.DataFrame(
        {
            "system_id": pd.Series(dtype="str"),
            "service_area": pd.Series(dtype="str"),
            "fresh_airflow_m3h": pd.Series(dtype="float64"),
            "outdoor_enthalpy_kjkg": pd.Series(dtype="float64"),
            "indoor_enthalpy_kjkg": pd.Series(dtype="float64"),
            "air_density_kgm3": pd.Series(dtype="float64"),
            "rated_cooling_kw": pd.Series(dtype="float64"),
            "rated_airflow_m3h": pd.Series(dtype="float64"),
        }
    )


def _as_dataframe(data):
    if isinstance(data, pd.DataFrame):
        return data.copy()
    return pd.DataFrame(data)


def _margin_status(value, upper_limit):
    if pd.isna(value):
        return "需复核"
    if value < 0:
        return "不足"
    if value <= upper_limit:
        return "合理"
    return "偏大"


def _safe_margin(numerator, denominator):
    if pd.isna(numerator) or pd.isna(denominator) or denominator <= 0:
        return pd.NA
    return numerator / denominator - 1


def _selection_conclusion(statuses):
    status_set = set(statuses)
    if "不足" in status_set:
        return "拟选偏小"
    if "需复核" in status_set or "焓差需复核" in status_set:
        return "需复核"
    if "偏大" in status_set:
        return "余量偏大"
    return "建议可用"


def calculate_fcu_selection(df):
    """计算风机盘管初步选型余量。"""
    result = _as_dataframe(df)
    result["冷量余量"] = result.apply(
        lambda row: _safe_margin(row["rated_cooling_kw"], row["cooling_load_kw"]),
        axis=1,
    )
    result["热量余量"] = result.apply(
        lambda row: _safe_margin(row["rated_heating_kw"], row["heating_load_kw"]),
        axis=1,
    )
    result["风量余量"] = result.apply(
        lambda row: _safe_margin(row["rated_airflow_m3h"], row["design_airflow_m3h"]),
        axis=1,
    )
    result["冷量校核"] = result["冷量余量"].apply(lambda v: _margin_status(v, 0.20))
    result["热量校核"] = result["热量余量"].apply(lambda v: _margin_status(v, 0.30))
    result["风量校核"] = result["风量余量"].apply(lambda v: _margin_status(v, 0.25))
    result["选型结论"] = result[["冷量校核", "热量校核", "风量校核"]].apply(
        _selection_conclusion,
        axis=1,
    )
    return result


def calculate_pau_selection(df):
    """计算新风机组初步选型余量。"""
    result = _as_dataframe(df)
    result["焓差 Δh (kJ/kg)"] = (
        result["outdoor_enthalpy_kjkg"] - result["indoor_enthalpy_kjkg"]
    )
    valid_enthalpy = result["焓差 Δh (kJ/kg)"] > 0
    result["新风质量流量 m (kg/s)"] = (
        result["air_density_kgm3"] * result["fresh_airflow_m3h"] / 3600
    )
    result["新风冷负荷 Q (kW)"] = pd.NA
    result.loc[valid_enthalpy, "新风冷负荷 Q (kW)"] = (
        result.loc[valid_enthalpy, "新风质量流量 m (kg/s)"]
        * result.loc[valid_enthalpy, "焓差 Δh (kJ/kg)"]
    )
    result["冷量余量"] = result.apply(
        lambda row: _safe_margin(row["rated_cooling_kw"], row["新风冷负荷 Q (kW)"]),
        axis=1,
    )
    result["风量余量"] = result.apply(
        lambda row: _safe_margin(row["rated_airflow_m3h"], row["fresh_airflow_m3h"]),
        axis=1,
    )
    result["焓差校核"] = result["焓差 Δh (kJ/kg)"].apply(
        lambda v: "适合估算" if pd.notna(v) and v > 0 else "焓差需复核"
    )
    result["冷量校核"] = result["冷量余量"].apply(lambda v: _margin_status(v, 0.20))
    result["风量校核"] = result["风量余量"].apply(lambda v: _margin_status(v, 0.25))
    result.loc[~valid_enthalpy, "冷量校核"] = "焓差需复核"
    result["选型结论"] = result[["焓差校核", "冷量校核", "风量校核"]].apply(
        _selection_conclusion,
        axis=1,
    )
    return result


def _rename_export_columns(df, col_map):
    available = {key: value for key, value in col_map.items() if key in df.columns}
    return df[list(available.keys())].rename(columns=available)


def _format_percent_columns(df, columns):
    formatted = df.copy()
    for col in columns:
        if col in formatted.columns:
            formatted[col] = formatted[col].apply(
                lambda v: "" if pd.isna(v) else f"{v:.1%}"
            )
    return formatted


def _build_terminal_word_report(title, description, df_export, summary_rows):
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(description)

    doc.add_heading("一、计算结果", level=1)
    table = doc.add_table(rows=1, cols=len(df_export.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df_export.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in df_export.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)

    doc.add_heading("二、汇总", level=1)
    for label, value in summary_rows:
        doc.add_paragraph(f"{label}：{value}")

    doc.add_heading("三、说明", level=1)
    doc.add_paragraph(
        "本说明书由建环工程计算工具箱自动生成，计算结果仅用于学习、课程设计辅助核算和工程初步校核。"
    )
    doc.add_paragraph(
        "设备最终选型应结合厂家样本、噪声、余压、水阻、控制方式和现行规范进行复核。"
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fcu_summary_rows(df):
    return [
        ("房间数量", f"{len(df)} 个"),
        ("冷量不足项", f"{(df['冷量校核'] == '不足').sum()} 项"),
        ("热量不足项", f"{(df['热量校核'] == '不足').sum()} 项"),
        ("风量不足项", f"{(df['风量校核'] == '不足').sum()} 项"),
    ]


def _pau_summary_rows(df):
    return [
        ("系统数量", f"{len(df)} 个"),
        ("冷量不足项", f"{(df['冷量校核'] == '不足').sum()} 项"),
        ("风量不足项", f"{(df['风量校核'] == '不足').sum()} 项"),
        ("焓差需复核项", f"{(df['焓差校核'] == '焓差需复核').sum()} 项"),
    ]


def _fcu_formula_table():
    st.markdown(
        FORMULA_STYLE
        + """
<table class="formula-table">
<thead>
<tr><th>计算项目</th><th>公式</th><th>单位说明</th></tr>
</thead>
<tbody>
<tr><td>冷量余量</td><td><span translate="no">η<sub>c</sub> = Q<sub>c,rated</sub> / Q<sub>c</sub> - 1</span></td><td><span translate="no">η<sub>c</sub></span>：冷量余量；<span translate="no">Q<sub>c,rated</sub></span>：额定冷量，<span translate="no">kW</span>；<span translate="no">Q<sub>c</sub></span>：房间冷负荷，<span translate="no">kW</span></td></tr>
<tr><td>热量余量</td><td><span translate="no">η<sub>h</sub> = Q<sub>h,rated</sub> / Q<sub>h</sub> - 1</span></td><td><span translate="no">η<sub>h</sub></span>：热量余量；<span translate="no">Q<sub>h,rated</sub></span>：额定热量，<span translate="no">kW</span>；<span translate="no">Q<sub>h</sub></span>：房间热负荷，<span translate="no">kW</span></td></tr>
<tr><td>风量余量</td><td><span translate="no">η<sub>L</sub> = L<sub>rated</sub> / L - 1</span></td><td><span translate="no">η<sub>L</sub></span>：风量余量；<span translate="no">L<sub>rated</sub></span>：额定风量，<span translate="no">m³/h</span>；<span translate="no">L</span>：设计风量，<span translate="no">m³/h</span></td></tr>
</tbody>
</table>
""",
        unsafe_allow_html=True,
    )


def _pau_formula_table():
    st.markdown(
        FORMULA_STYLE
        + """
<table class="formula-table">
<thead>
<tr><th>计算项目</th><th>公式</th><th>单位说明</th></tr>
</thead>
<tbody>
<tr><td>新风质量流量</td><td><span translate="no">m = ρ × L / 3600</span></td><td><span translate="no">m</span>：<span translate="no">kg/s</span>；ρ：<span translate="no">kg/m³</span>；<span translate="no">L</span>：<span translate="no">m³/h</span></td></tr>
<tr><td>新风冷负荷</td><td><span translate="no">Q = m × (h<sub>W</sub> - h<sub>N</sub>)</span></td><td><span translate="no">Q</span>：<span translate="no">kW</span>；<span translate="no">h<sub>W</sub></span>、<span translate="no">h<sub>N</sub></span>：<span translate="no">kJ/kg</span></td></tr>
<tr><td>机组冷量余量</td><td><span translate="no">η<sub>Q</sub> = Q<sub>rated</sub> / Q - 1</span></td><td><span translate="no">η<sub>Q</sub></span>：冷量余量；<span translate="no">Q<sub>rated</sub></span>：拟选机组冷量，<span translate="no">kW</span></td></tr>
<tr><td>机组风量余量</td><td><span translate="no">η<sub>L</sub> = L<sub>rated</sub> / L - 1</span></td><td><span translate="no">η<sub>L</sub></span>：风量余量；<span translate="no">L<sub>rated</sub></span>：拟选机组风量，<span translate="no">m³/h</span></td></tr>
</tbody>
</table>
""",
        unsafe_allow_html=True,
    )


def _render_fcu_tab():
    st.subheader("风机盘管初步选型")
    st.caption(
        "余量阈值为简化建议，可按项目要求调整，不能替代设备样本和规范校核。"
    )

    if "terminal_fcu_df" not in st.session_state:
        st.session_state["terminal_fcu_df"] = _empty_fcu_df()
    if "terminal_fcu_editor_version" not in st.session_state:
        st.session_state["terminal_fcu_editor_version"] = 0

    col_load, col_calc, col_clear = st.columns([2, 1, 1])
    with col_load:
        if st.button("加载风机盘管示例数据", use_container_width=True):
            st.session_state["terminal_fcu_df"] = pd.DataFrame(FCU_SAMPLE_DATA)
            st.session_state["terminal_fcu_editor_version"] += 1
            st.success("已加载风机盘管示例数据。")
    with col_calc:
        st.button("计算", use_container_width=True, type="primary")
    with col_clear:
        if st.button("清空", key="clear_fcu", use_container_width=True):
            st.session_state["terminal_fcu_df"] = _empty_fcu_df()
            st.session_state["terminal_fcu_editor_version"] += 1

    edited_df = st.data_editor(
        st.session_state["terminal_fcu_df"],
        num_rows="dynamic",
        use_container_width=True,
        key=f"terminal_fcu_editor_{st.session_state['terminal_fcu_editor_version']}",
        column_config={
            "room_id": st.column_config.TextColumn("房间编号"),
            "room_name": st.column_config.TextColumn("房间名称"),
            "cooling_load_kw": st.column_config.NumberColumn("冷负荷 Qc (kW)", min_value=0, format="%.2f"),
            "heating_load_kw": st.column_config.NumberColumn("热负荷 Qh (kW)", min_value=0, format="%.2f"),
            "design_airflow_m3h": st.column_config.NumberColumn("设计风量 L (m³/h)", min_value=0, format="%.0f"),
            "rated_cooling_kw": st.column_config.NumberColumn("额定冷量 (kW)", min_value=0, format="%.2f"),
            "rated_heating_kw": st.column_config.NumberColumn("额定热量 (kW)", min_value=0, format="%.2f"),
            "rated_airflow_m3h": st.column_config.NumberColumn("额定风量 (m³/h)", min_value=0, format="%.0f"),
        },
    )
    st.session_state["terminal_fcu_df"] = edited_df

    st.subheader("计算结果")
    if edited_df.empty:
        st.info("请输入房间数据或加载示例数据开始计算。")
        return

    required = [
        "cooling_load_kw",
        "heating_load_kw",
        "design_airflow_m3h",
        "rated_cooling_kw",
        "rated_heating_kw",
        "rated_airflow_m3h",
    ]
    valid = (edited_df[required].fillna(0) > 0).all(axis=1)
    if (~valid).any():
        st.warning("存在零值、负值或缺失值的房间，相关余量会标记为需复核。")

    result = calculate_fcu_selection(edited_df)
    display_cols = [
        "room_id",
        "room_name",
        "cooling_load_kw",
        "heating_load_kw",
        "design_airflow_m3h",
        "rated_cooling_kw",
        "rated_heating_kw",
        "rated_airflow_m3h",
        "冷量余量",
        "热量余量",
        "风量余量",
        "冷量校核",
        "热量校核",
        "风量校核",
        "选型结论",
    ]
    df_display = result[display_cols].copy()
    for col in ["冷量余量", "热量余量", "风量余量"]:
        df_display[col] = df_display[col].apply(lambda v: None if pd.isna(v) else v)
    st.dataframe(
        df_display,
        use_container_width=True,
        hide_index=True,
        column_config={
            "room_id": "房间编号",
            "room_name": "房间名称",
            "cooling_load_kw": st.column_config.NumberColumn("冷负荷 Qc (kW)", format="%.2f"),
            "heating_load_kw": st.column_config.NumberColumn("热负荷 Qh (kW)", format="%.2f"),
            "design_airflow_m3h": st.column_config.NumberColumn("设计风量 L (m³/h)", format="%.0f"),
            "rated_cooling_kw": st.column_config.NumberColumn("额定冷量 (kW)", format="%.2f"),
            "rated_heating_kw": st.column_config.NumberColumn("额定热量 (kW)", format="%.2f"),
            "rated_airflow_m3h": st.column_config.NumberColumn("额定风量 (m³/h)", format="%.0f"),
            "冷量余量": st.column_config.NumberColumn("冷量余量", format="%.1%"),
            "热量余量": st.column_config.NumberColumn("热量余量", format="%.1%"),
            "风量余量": st.column_config.NumberColumn("风量余量", format="%.1%"),
        },
    )

    st.divider()
    st.subheader("系统汇总")
    cols = st.columns(4)
    for col_obj, (label, value) in zip(cols, _fcu_summary_rows(result)):
        col_obj.metric(label, value)

    st.divider()
    st.subheader("导出数据")
    df_export = _format_percent_columns(_rename_export_columns(result, FCU_EXPORT_MAP), ["冷量余量", "热量余量", "风量余量"])
    summary_rows = _fcu_summary_rows(result)
    col_csv, col_xlsx = st.columns(2)
    with col_csv:
        st.download_button(
            label="📥 下载 CSV",
            data=get_csv_bytes(df_export),
            file_name="风机盘管初步选型结果.csv",
            mime="text/csv",
            use_container_width=True,
        )
    with col_xlsx:
        st.download_button(
            label="📥 下载 Excel",
            data=export_formatted_excel(df_export, summary_rows, sheet_name="风机盘管选型", summary_sheet_name="选型汇总"),
            file_name="风机盘管初步选型结果.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    st.download_button(
        label="📄 导出 Word 计算说明书",
        data=_build_terminal_word_report(
            "风机盘管初步选型计算说明书",
            "用于风机盘管冷量、热量和风量的简化初步选型校核。",
            df_export,
            summary_rows,
        ),
        file_name="风机盘管初步选型计算说明书.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.divider()
    st.subheader("计算公式说明")
    with st.expander("查看风机盘管计算公式"):
        _fcu_formula_table()


def _render_pau_tab():
    st.subheader("新风机组初步选型")
    st.caption(
        "当室外空气焓值不高于室内空气焓值时，不用于夏季新风冷负荷估算，相关冷量结果标记为需复核。"
    )

    if "terminal_pau_df" not in st.session_state:
        st.session_state["terminal_pau_df"] = _empty_pau_df()
    if "terminal_pau_editor_version" not in st.session_state:
        st.session_state["terminal_pau_editor_version"] = 0

    col_load, col_calc, col_clear = st.columns([2, 1, 1])
    with col_load:
        if st.button("加载新风机组示例数据", use_container_width=True):
            st.session_state["terminal_pau_df"] = pd.DataFrame(PAU_SAMPLE_DATA)
            st.session_state["terminal_pau_editor_version"] += 1
            st.success("已加载新风机组示例数据。")
    with col_calc:
        st.button("计算", key="calc_pau", use_container_width=True, type="primary")
    with col_clear:
        if st.button("清空", key="clear_pau", use_container_width=True):
            st.session_state["terminal_pau_df"] = _empty_pau_df()
            st.session_state["terminal_pau_editor_version"] += 1

    edited_df = st.data_editor(
        st.session_state["terminal_pau_df"],
        num_rows="dynamic",
        use_container_width=True,
        key=f"terminal_pau_editor_{st.session_state['terminal_pau_editor_version']}",
        column_config={
            "system_id": st.column_config.TextColumn("系统编号"),
            "service_area": st.column_config.TextColumn("服务区域"),
            "fresh_airflow_m3h": st.column_config.NumberColumn("新风量 L (m³/h)", min_value=0, format="%.0f"),
            "outdoor_enthalpy_kjkg": st.column_config.NumberColumn("室外空气焓值 hW (kJ/kg)", format="%.1f"),
            "indoor_enthalpy_kjkg": st.column_config.NumberColumn("室内空气焓值 hN (kJ/kg)", format="%.1f"),
            "air_density_kgm3": st.column_config.NumberColumn("空气密度 ρ (kg/m³)", min_value=0, format="%.2f"),
            "rated_cooling_kw": st.column_config.NumberColumn("拟选新风机组冷量 (kW)", min_value=0, format="%.2f"),
            "rated_airflow_m3h": st.column_config.NumberColumn("拟选新风机组风量 (m³/h)", min_value=0, format="%.0f"),
        },
    )
    st.session_state["terminal_pau_df"] = edited_df

    st.subheader("计算结果")
    if edited_df.empty:
        st.info("请输入新风系统数据或加载示例数据开始计算。")
        return

    result = calculate_pau_selection(edited_df)
    if (result["焓差校核"] == "焓差需复核").any():
        st.warning("存在 hW - hN <= 0 的系统，已避免显示负冷负荷，请复核工况和焓值。")

    display_cols = [
        "system_id",
        "service_area",
        "fresh_airflow_m3h",
        "outdoor_enthalpy_kjkg",
        "indoor_enthalpy_kjkg",
        "air_density_kgm3",
        "新风质量流量 m (kg/s)",
        "新风冷负荷 Q (kW)",
        "rated_cooling_kw",
        "rated_airflow_m3h",
        "冷量余量",
        "风量余量",
        "冷量校核",
        "风量校核",
        "焓差校核",
        "选型结论",
    ]
    df_display = result[display_cols].copy()
    for col in ["冷量余量", "风量余量", "新风冷负荷 Q (kW)"]:
        df_display[col] = df_display[col].apply(lambda v: None if pd.isna(v) else v)
    st.dataframe(
        df_display,
        use_container_width=True,
        hide_index=True,
        column_config={
            "system_id": "系统编号",
            "service_area": "服务区域",
            "fresh_airflow_m3h": st.column_config.NumberColumn("新风量 L (m³/h)", format="%.0f"),
            "outdoor_enthalpy_kjkg": st.column_config.NumberColumn("室外空气焓值 hW (kJ/kg)", format="%.1f"),
            "indoor_enthalpy_kjkg": st.column_config.NumberColumn("室内空气焓值 hN (kJ/kg)", format="%.1f"),
            "air_density_kgm3": st.column_config.NumberColumn("空气密度 ρ (kg/m³)", format="%.2f"),
            "新风质量流量 m (kg/s)": st.column_config.NumberColumn("新风质量流量 m (kg/s)", format="%.3f"),
            "新风冷负荷 Q (kW)": st.column_config.NumberColumn("新风冷负荷 Q (kW)", format="%.2f"),
            "rated_cooling_kw": st.column_config.NumberColumn("拟选新风机组冷量 (kW)", format="%.2f"),
            "rated_airflow_m3h": st.column_config.NumberColumn("拟选新风机组风量 (m³/h)", format="%.0f"),
            "冷量余量": st.column_config.NumberColumn("冷量余量", format="%.1%"),
            "风量余量": st.column_config.NumberColumn("风量余量", format="%.1%"),
        },
    )

    st.divider()
    st.subheader("系统汇总")
    cols = st.columns(4)
    for col_obj, (label, value) in zip(cols, _pau_summary_rows(result)):
        col_obj.metric(label, value)

    st.divider()
    st.subheader("导出数据")
    df_export = _format_percent_columns(_rename_export_columns(result, PAU_EXPORT_MAP), ["冷量余量", "风量余量"])
    summary_rows = _pau_summary_rows(result)
    col_csv, col_xlsx = st.columns(2)
    with col_csv:
        st.download_button(
            label="📥 下载 CSV",
            data=get_csv_bytes(df_export),
            file_name="新风机组初步选型结果.csv",
            mime="text/csv",
            use_container_width=True,
        )
    with col_xlsx:
        st.download_button(
            label="📥 下载 Excel",
            data=export_formatted_excel(df_export, summary_rows, sheet_name="新风机组选型", summary_sheet_name="选型汇总"),
            file_name="新风机组初步选型结果.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    st.download_button(
        label="📄 导出 Word 计算说明书",
        data=_build_terminal_word_report(
            "新风机组初步选型计算说明书",
            "用于新风机组冷量和风量的简化初步选型校核。",
            df_export,
            summary_rows,
        ),
        file_name="新风机组初步选型计算说明书.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.divider()
    st.subheader("计算公式说明")
    with st.expander("查看新风机组计算公式"):
        _pau_formula_table()


def render_terminal_equipment_module():
    """渲染空调末端设备初步选型模块。"""

    st.markdown("### 模块三：空调末端设备初步选型")
    st.markdown(
        "本模块用于风机盘管和新风机组的初步选型校核，可根据房间冷负荷、热负荷、"
        "新风量、空气焓值等参数，估算所需设备冷量、热量和风量，并给出选型余量建议。"
        "计算结果仅用于学习、课程设计辅助核算和工程初步校核，不能替代正式设备样本选型。"
    )

    fcu_tab, pau_tab = st.tabs(["风机盘管初步选型", "新风机组初步选型"])
    with fcu_tab:
        _render_fcu_tab()
    with pau_tab:
        _render_pau_tab()

    st.divider()
    st.caption(
        "末端设备选型结果仅供学习、课程设计辅助核算和工程初步校核使用，"
        "实际工程应结合设备样本、噪声、余压、水阻、控制方式和规范要求复核。"
    )
